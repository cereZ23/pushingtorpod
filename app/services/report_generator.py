"""
Report generator orchestrator for EASM PDF and DOCX reports.

Collects data from the database (reusing the same queries as the JSON
endpoints in ``app.api.routers.reports``), generates charts via
``chart_generator``, and renders the final document using either
WeasyPrint (PDF) or docxtpl (DOCX).
"""

import io
import json
import logging
import os
import uuid
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from jinja2 import Environment, FileSystemLoader
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.models.database import (
    Asset,
    AssetType,
    Finding,
    FindingSeverity,
    FindingStatus,
    Tenant,
)
from app.models.risk import RiskScore
from app.services.chart_generator import (
    chart_to_data_uri,
    generate_asset_chart,
    generate_risk_gauge,
    generate_severity_chart,
    generate_trend_chart,
)

logger = logging.getLogger(__name__)

# Path to template directory
TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates" / "reports"

# Severity weights (same as reports.py)
SEVERITY_WEIGHT: Dict[str, int] = {
    "critical": 50,
    "high": 30,
    "medium": 15,
    "low": 5,
    "info": 1,
}

# Remediation guidance mapping
_REMEDIATION_MAP: Dict[str, str] = {
    "http-missing-security-headers": "Add the missing HTTP security header(s) in the web server configuration.",
    "ssl-detect": "Ensure TLS 1.2+ is enforced and weak cipher suites are disabled.",
    "exposed-panels": "Restrict access to administrative panels via IP allowlisting or VPN.",
    "default-login": "Change default credentials and enforce strong password policies.",
}


# ---------------------------------------------------------------------------
# Compliance framework mapping
# ---------------------------------------------------------------------------


class _FrameworkCategory:
    """Lightweight container for a compliance framework category."""

    __slots__ = ("id", "name", "description")

    def __init__(self, id: str, name: str, description: str = ""):
        self.id = id
        self.name = name
        self.description = description


# SOC 2 Trust Service Criteria categories
TSC_CATEGORIES: List[_FrameworkCategory] = [
    _FrameworkCategory(
        "CC1",
        "Control Environment",
        "Management's commitment to integrity and ethical values, governance oversight, and organizational structure.",
    ),
    _FrameworkCategory(
        "CC2",
        "Communication and Information",
        "Information needed to carry out internal control responsibilities and communication of objectives.",
    ),
    _FrameworkCategory(
        "CC3",
        "Risk Assessment",
        "Identification and assessment of risks to the achievement of the entity's objectives.",
    ),
    _FrameworkCategory(
        "CC4",
        "Monitoring Activities",
        "Selection, development, and performance of ongoing and/or separate evaluations.",
    ),
    _FrameworkCategory(
        "CC5", "Control Activities", "Selection and development of control activities that mitigate risks."
    ),
    _FrameworkCategory(
        "CC6",
        "Logical and Physical Access Controls",
        "Logical access security, physical access restrictions, and data protection.",
    ),
    _FrameworkCategory(
        "CC7",
        "System Operations",
        "Detection and monitoring of anomalies, including security incidents and vulnerabilities.",
    ),
    _FrameworkCategory(
        "CC8", "Change Management", "Authorized, tested, and approved system changes, including emergency changes."
    ),
    _FrameworkCategory(
        "CC9", "Risk Mitigation", "Risk mitigation through business continuity and vendor management activities."
    ),
]

# ISO 27001:2022 Annex A control domains
ANNEX_A_DOMAINS: List[_FrameworkCategory] = [
    _FrameworkCategory(
        "A.5",
        "Information Security Policies",
        "Management direction for information security in accordance with business requirements and relevant laws.",
    ),
    _FrameworkCategory(
        "A.6",
        "Organization of Information Security",
        "Internal organization, mobile devices, and teleworking security controls.",
    ),
    _FrameworkCategory(
        "A.7", "Human Resource Security", "Security controls for before, during, and termination of employment."
    ),
    _FrameworkCategory(
        "A.8", "Asset Management", "Responsibility for assets, information classification, and media handling."
    ),
    _FrameworkCategory(
        "A.9", "Access Control", "Business requirements, user access management, and system/application access control."
    ),
    _FrameworkCategory(
        "A.10",
        "Cryptography",
        "Cryptographic controls for protection of information confidentiality, integrity, and authenticity.",
    ),
    _FrameworkCategory(
        "A.11", "Physical and Environmental Security", "Secure areas, equipment security, and environmental controls."
    ),
    _FrameworkCategory(
        "A.12",
        "Operations Security",
        "Operational procedures, malware protection, backup, logging, and technical vulnerability management.",
    ),
    _FrameworkCategory(
        "A.13", "Communications Security", "Network security management, information transfer policies, and controls."
    ),
    _FrameworkCategory(
        "A.14",
        "System Acquisition, Development and Maintenance",
        "Security requirements, development/support processes, and test data protection.",
    ),
    _FrameworkCategory(
        "A.15",
        "Supplier Relationships",
        "Information security in supplier relationships and service delivery management.",
    ),
    _FrameworkCategory(
        "A.16",
        "Information Security Incident Management",
        "Management of incidents, reporting, and collection of evidence.",
    ),
    _FrameworkCategory("A.17", "Business Continuity Management", "Information security continuity and redundancies."),
    _FrameworkCategory(
        "A.18",
        "Compliance",
        "Compliance with legal, contractual, and policy requirements; information security reviews.",
    ),
]

# Mapping: Nuclei template prefix → (SOC2 TSC ID, ISO 27001 Annex A ID)
# This maps common EASM finding types to the most relevant compliance controls.
_TEMPLATE_COMPLIANCE_MAP: Dict[str, Dict[str, str]] = {
    # TLS/SSL findings → Cryptography / Logical Access
    "ssl-": {"tsc": "CC6", "annex": "A.10"},
    "tls-": {"tsc": "CC6", "annex": "A.10"},
    "weak-cipher": {"tsc": "CC6", "annex": "A.10"},
    "expired-ssl": {"tsc": "CC6", "annex": "A.10"},
    # HTTP security headers → System Operations / Communications
    "http-missing-security-headers": {"tsc": "CC7", "annex": "A.14"},
    "strict-transport-security": {"tsc": "CC6", "annex": "A.10"},
    "content-security-policy": {"tsc": "CC7", "annex": "A.14"},
    "x-frame-options": {"tsc": "CC7", "annex": "A.14"},
    "x-content-type": {"tsc": "CC7", "annex": "A.14"},
    "permissions-policy": {"tsc": "CC7", "annex": "A.14"},
    # Exposed panels / Default credentials → Access Control
    "exposed-panels": {"tsc": "CC6", "annex": "A.9"},
    "default-login": {"tsc": "CC6", "annex": "A.9"},
    "default-credentials": {"tsc": "CC6", "annex": "A.9"},
    "admin-panel": {"tsc": "CC6", "annex": "A.9"},
    "login-panel": {"tsc": "CC6", "annex": "A.9"},
    # Misconfigurations → Change Management / Operations
    "misconfiguration": {"tsc": "CC8", "annex": "A.12"},
    "misconfig": {"tsc": "CC8", "annex": "A.12"},
    "debug-enabled": {"tsc": "CC8", "annex": "A.12"},
    "directory-listing": {"tsc": "CC6", "annex": "A.12"},
    "cors-misconfig": {"tsc": "CC7", "annex": "A.14"},
    # CVEs / Known vulnerabilities → Risk Assessment / Vulnerability Management
    "cve-": {"tsc": "CC3", "annex": "A.12"},
    "CVE-": {"tsc": "CC3", "annex": "A.12"},
    # Information disclosure → Monitoring
    "information-disclosure": {"tsc": "CC7", "annex": "A.12"},
    "tech-detect": {"tsc": "CC7", "annex": "A.8"},
    "wappalyzer": {"tsc": "CC7", "annex": "A.8"},
    # DNS / Network findings → Communications Security
    "dns-": {"tsc": "CC7", "annex": "A.13"},
    "zone-transfer": {"tsc": "CC6", "annex": "A.13"},
    "subdomain-takeover": {"tsc": "CC6", "annex": "A.13"},
    "open-redirect": {"tsc": "CC7", "annex": "A.14"},
    # Cloud / Infrastructure → Risk Mitigation
    "cloud-": {"tsc": "CC9", "annex": "A.15"},
    "s3-bucket": {"tsc": "CC6", "annex": "A.9"},
    "azure-": {"tsc": "CC9", "annex": "A.15"},
    "gcp-": {"tsc": "CC9", "annex": "A.15"},
    # Email / DMARC / SPF → Communications
    "dmarc": {"tsc": "CC7", "annex": "A.13"},
    "spf": {"tsc": "CC7", "annex": "A.13"},
    "dkim": {"tsc": "CC7", "annex": "A.13"},
    # Sensitive paths / Files → Access Control
    "sensitive-": {"tsc": "CC6", "annex": "A.9"},
    "backup-file": {"tsc": "CC6", "annex": "A.9"},
    "git-config": {"tsc": "CC6", "annex": "A.9"},
    "env-file": {"tsc": "CC6", "annex": "A.9"},
    ".env": {"tsc": "CC6", "annex": "A.9"},
}

# Severity-based fallback mapping for findings without template match
_SEVERITY_COMPLIANCE_FALLBACK: Dict[str, Dict[str, str]] = {
    "critical": {"tsc": "CC3", "annex": "A.12"},
    "high": {"tsc": "CC7", "annex": "A.12"},
    "medium": {"tsc": "CC7", "annex": "A.14"},
    "low": {"tsc": "CC4", "annex": "A.12"},
    "info": {"tsc": "CC4", "annex": "A.8"},
}


def _map_finding_to_framework(
    template_id: Optional[str],
    name: str,
    severity: str,
    framework: str,
) -> str:
    """Map a finding to a compliance framework control ID.

    Tries template_id prefix matching first, then falls back to name-based
    matching, then severity-based default.

    Args:
        template_id: Nuclei template ID.
        name: Finding name.
        severity: Finding severity string.
        framework: ``"tsc"`` for SOC 2 or ``"annex"`` for ISO 27001.

    Returns:
        Control ID string (e.g. ``"CC6"`` or ``"A.12"``).
    """
    # Try template_id prefix match
    if template_id:
        for prefix, mapping in _TEMPLATE_COMPLIANCE_MAP.items():
            if template_id.startswith(prefix) or prefix in template_id:
                return mapping[framework]

    # Try name-based match
    name_lower = name.lower() if name else ""
    for prefix, mapping in _TEMPLATE_COMPLIANCE_MAP.items():
        if prefix.lower().rstrip("-") in name_lower:
            return mapping[framework]

    # Fallback to severity-based mapping
    fallback = _SEVERITY_COMPLIANCE_FALLBACK.get(severity, _SEVERITY_COMPLIANCE_FALLBACK["info"])
    return fallback[framework]


def _get_remediation(template_id: Optional[str]) -> Optional[str]:
    if not template_id:
        return None
    for prefix, guidance in _REMEDIATION_MAP.items():
        if template_id.startswith(prefix):
            return guidance
    return None


def _format_evidence(evidence) -> str:
    """Convert evidence (str or dict) to a readable text block for PDF."""
    if evidence is None:
        return ""
    if isinstance(evidence, dict):
        return json.dumps(evidence, indent=2, default=str)[:500]
    if isinstance(evidence, str):
        try:
            parsed = json.loads(evidence)
            if isinstance(parsed, dict):
                return json.dumps(parsed, indent=2, default=str)[:500]
            return str(parsed)[:500]
        except (json.JSONDecodeError, TypeError):
            return evidence[:500]
    return str(evidence)[:500]


class ReportGenerator:
    """Generates PDF and DOCX reports for a tenant."""

    def __init__(self, db: Session, tenant_id: int):
        self.db = db
        self.tenant_id = tenant_id
        self._tenant: Optional[Tenant] = None

    @property
    def tenant(self) -> Tenant:
        if self._tenant is None:
            self._tenant = self.db.query(Tenant).filter(Tenant.id == self.tenant_id).first()
        return self._tenant

    # ------------------------------------------------------------------
    # Data collection (mirrors reports.py logic)
    # ------------------------------------------------------------------

    def _collect_executive_data(self) -> Dict[str, Any]:
        """Gather all data needed for the executive report."""
        db = self.db
        tid = self.tenant_id

        # Risk score
        latest_risk = (
            db.query(RiskScore)
            .filter(RiskScore.tenant_id == tid, RiskScore.scope_type == "organization")
            .order_by(RiskScore.scored_at.desc())
            .first()
        )
        risk_score = round(latest_risk.score, 2) if latest_risk else 0.0
        risk_grade = latest_risk.grade if latest_risk else "N/A"

        # Score trend
        cutoff = datetime.now(timezone.utc) - timedelta(days=30)
        score_rows = (
            db.query(RiskScore)
            .filter(
                RiskScore.tenant_id == tid,
                RiskScore.scope_type == "organization",
                RiskScore.scored_at >= cutoff,
            )
            .order_by(RiskScore.scored_at.asc())
            .all()
        )
        score_trend = [
            {"date": row.scored_at, "score": round(row.score, 2), "grade": row.grade or "N/A"} for row in score_rows
        ]

        # Asset counts
        asset_counts: Dict[str, int] = {}
        for at in AssetType:
            c = db.query(func.count(Asset.id)).filter(Asset.tenant_id == tid, Asset.type == at).scalar() or 0
            asset_counts[at.value] = c
        total_assets = sum(asset_counts.values())

        # Finding counts by severity
        sev_counts: Dict[str, int] = {}
        for sev in FindingSeverity:
            c = (
                db.query(func.count(Finding.id))
                .join(Asset)
                .filter(Asset.tenant_id == tid, Finding.severity == sev)
                .scalar()
                or 0
            )
            sev_counts[sev.value] = c

        # Finding counts by status
        status_counts: Dict[str, int] = {}
        for st in FindingStatus:
            c = (
                db.query(func.count(Finding.id))
                .join(Asset)
                .filter(Asset.tenant_id == tid, Finding.status == st)
                .scalar()
                or 0
            )
            status_counts[st.value] = c

        total_findings = sum(status_counts.values())
        open_findings = status_counts.get("open", 0)

        # Top 10 issues
        top_query = (
            db.query(Finding, Asset.identifier)
            .join(Asset)
            .filter(Asset.tenant_id == tid, Finding.status == FindingStatus.OPEN)
            .order_by(Finding.severity.desc(), Finding.cvss_score.desc().nullslast())
            .limit(50)
            .all()
        )
        scored = []
        for finding, asset_id in top_query:
            w = SEVERITY_WEIGHT.get(finding.severity.value, 0)
            cvss = finding.cvss_score or 0.0
            scored.append((w + cvss, finding, asset_id))
        scored.sort(key=lambda x: x[0], reverse=True)

        top_issues = [
            {
                "finding_id": f.id,
                "name": f.name,
                "severity": f.severity.value,
                "cvss_score": f.cvss_score,
                "cve_id": f.cve_id,
                "template_id": f.template_id,
                "asset_identifier": aid,
                "first_seen": f.first_seen.strftime("%Y-%m-%d") if f.first_seen else "-",
            }
            for _, f, aid in scored[:10]
        ]

        # Recommendations
        recommendations = self._build_recommendations(sev_counts, open_findings)

        return {
            "risk_score": risk_score,
            "risk_grade": risk_grade,
            "score_trend": score_trend,
            "asset_counts": asset_counts,
            "severity_counts": sev_counts,
            "total_assets": total_assets,
            "total_findings": total_findings,
            "open_findings": open_findings,
            "top_issues": top_issues,
            "recommendations": recommendations,
        }

    def _collect_findings(
        self,
        severity: Optional[str] = None,
        finding_status: Optional[str] = None,
        limit: int = 500,
    ) -> List[Dict[str, Any]]:
        """Collect detailed findings for the technical report."""
        db = self.db
        query = (
            db.query(Finding, Asset.identifier, Asset.type)
            .join(Asset)
            .filter(
                Asset.tenant_id == self.tenant_id,
                Asset.is_active.is_(True),
            )
        )

        if severity:
            try:
                query = query.filter(Finding.severity == FindingSeverity(severity))
            except ValueError:
                pass

        if finding_status:
            try:
                query = query.filter(Finding.status == FindingStatus(finding_status))
            except ValueError:
                pass
        else:
            # Default: only open findings (exclude fixed/suppressed)
            query = query.filter(Finding.status == FindingStatus.OPEN)

        results = (
            query.order_by(
                Finding.severity.desc(),
                Finding.cvss_score.desc().nullslast(),
                Finding.last_seen.desc(),
            )
            .limit(limit)
            .all()
        )

        return [
            {
                "id": f.id,
                "name": f.name,
                "severity": f.severity.value,
                "cvss_score": f.cvss_score,
                "cve_id": f.cve_id,
                "template_id": f.template_id,
                "source": f.source,
                "status": f.status.value,
                "asset_identifier": aid,
                "asset_type": atype.value if atype else "unknown",
                "evidence": f.evidence,
                "first_seen": f.first_seen.strftime("%Y-%m-%d") if f.first_seen else "-",
                "last_seen": f.last_seen.strftime("%Y-%m-%d") if f.last_seen else "-",
                "occurrence_count": getattr(f, "occurrence_count", 1) or 1,
                "remediation": _get_remediation(f.template_id),
            }
            for f, aid, atype in results
        ]

    @staticmethod
    def _build_recommendations(sev_counts: Dict[str, int], open_count: int) -> List[Dict[str, Any]]:
        recs: List[Dict[str, Any]] = []
        priority = 1

        critical = sev_counts.get("critical", 0)
        high = sev_counts.get("high", 0)
        medium = sev_counts.get("medium", 0)

        if critical > 0:
            recs.append(
                {
                    "priority": priority,
                    "title": "Remediate critical vulnerabilities immediately",
                    "description": (
                        f"{critical} critical finding(s) detected. These represent the highest risk "
                        "and should be patched or mitigated within 24 hours."
                    ),
                    "affected_count": critical,
                }
            )
            priority += 1

        if high > 0:
            recs.append(
                {
                    "priority": priority,
                    "title": "Address high-severity findings within SLA",
                    "description": (
                        f"{high} high-severity finding(s) require attention. Schedule remediation "
                        "within the 7-day SLA window."
                    ),
                    "affected_count": high,
                }
            )
            priority += 1

        if medium > 0:
            recs.append(
                {
                    "priority": priority,
                    "title": "Plan remediation of medium-severity issues",
                    "description": (
                        f"{medium} medium-severity finding(s) should be addressed in the next "
                        "sprint or maintenance cycle."
                    ),
                    "affected_count": medium,
                }
            )
            priority += 1

        if open_count > 0:
            recs.append(
                {
                    "priority": priority,
                    "title": "Review and triage all open findings",
                    "description": (
                        f"{open_count} finding(s) are currently open. Triage to suppress "
                        "false positives and assign ownership."
                    ),
                    "affected_count": open_count,
                }
            )
            priority += 1

        recs.append(
            {
                "priority": priority,
                "title": "Maintain continuous monitoring",
                "description": (
                    "Ensure scheduled scans are active and alert policies configured "
                    "for critical and high-severity events."
                ),
                "affected_count": 0,
            }
        )

        return recs

    # ------------------------------------------------------------------
    # Compliance report data collection
    # ------------------------------------------------------------------

    def _collect_compliance_data(self, framework: str) -> Dict[str, Any]:
        """Collect data for a compliance report (SOC 2 or ISO 27001).

        Reuses executive data as a base, then maps every open finding to the
        relevant framework control and builds the compliance-specific context
        expected by the Jinja2 templates.

        Args:
            framework: ``"soc2"`` or ``"iso27001"``.

        Returns:
            Template context dict with compliance mapping, gap analysis,
            non-conformity lists, and all shared executive data.
        """
        base = self._collect_executive_data()

        # Fetch all findings (open + suppressed, excluding fixed)
        findings = self._collect_findings(finding_status=None, limit=5000)

        # Determine framework-specific structures
        if framework == "soc2":
            categories = TSC_CATEGORIES
            fw_key = "tsc"
            ctrl_id_key = "tsc_id"
        else:
            categories = ANNEX_A_DOMAINS
            fw_key = "annex"
            ctrl_id_key = "annex_id"

        # Map findings to controls
        compliance_map: Dict[str, Dict[str, Any]] = {}
        for cat in categories:
            compliance_map[cat.id] = {"findings": []}

        for f in findings:
            ctrl = _map_finding_to_framework(
                template_id=f.get("template_id"),
                name=f.get("name", ""),
                severity=f.get("severity", "info"),
                framework=fw_key,
            )
            f[ctrl_id_key] = ctrl
            if ctrl in compliance_map:
                compliance_map[ctrl]["findings"].append(f)

        # Gap analysis (SOC 2) / Non-conformities (ISO 27001)
        control_failures = [f for f in findings if f.get("severity") == "critical"]
        control_weaknesses = [f for f in findings if f.get("severity") == "high"]
        observations = [f for f in findings if f.get("severity") == "medium"]

        # Count controls with issues
        gap_count = sum(
            1
            for cat_data in compliance_map.values()
            if any(f.get("severity") in ("critical", "high", "medium") for f in cat_data["findings"])
        )

        context = {
            **base,
            "compliance_map": compliance_map,
        }

        if framework == "soc2":
            context.update(
                {
                    "tsc_categories": categories,
                    "tsc_total_controls": len(categories),
                    "tsc_gap_count": gap_count,
                    "gap_analysis": {
                        "control_failures": control_failures,
                        "control_weaknesses": control_weaknesses,
                        "observations": observations,
                    },
                }
            )
        else:
            context.update(
                {
                    "annex_a_domains": categories,
                    "annex_a_total_controls": len(categories),
                    "annex_a_nc_count": gap_count,
                    "nonconformities": {
                        "major": control_failures,
                        "minor": control_weaknesses,
                        "observations": observations,
                    },
                }
            )

        return context

    # ------------------------------------------------------------------
    # Chart generation
    # ------------------------------------------------------------------

    def _generate_charts(self, data: Dict[str, Any], fmt: str = "svg") -> Dict[str, str]:
        """Generate all charts and return data URIs."""
        charts: Dict[str, str] = {}

        gauge_bytes = generate_risk_gauge(data["risk_score"], data["risk_grade"], fmt=fmt)
        charts["gauge_chart"] = chart_to_data_uri(gauge_bytes, fmt)

        sev_bytes = generate_severity_chart(data["severity_counts"], fmt=fmt)
        charts["severity_chart"] = chart_to_data_uri(sev_bytes, fmt)

        asset_bytes = generate_asset_chart(data["asset_counts"], fmt=fmt)
        charts["asset_chart"] = chart_to_data_uri(asset_bytes, fmt)

        trend_bytes = generate_trend_chart(data["score_trend"], fmt=fmt)
        if trend_bytes:
            charts["trend_chart"] = chart_to_data_uri(trend_bytes, fmt)
        else:
            charts["trend_chart"] = ""

        return charts

    # ------------------------------------------------------------------
    # PDF generation
    # ------------------------------------------------------------------

    def generate_pdf(
        self,
        report_type: str = "executive",
        severity: Optional[str] = None,
        finding_status: Optional[str] = None,
        limit: int = 500,
    ) -> bytes:
        """
        Generate a PDF report.

        Args:
            report_type: ``"executive"``, ``"technical"``, ``"soc2"``,
                         or ``"iso27001"``.
            severity: Optional severity filter (technical only).
            finding_status: Optional status filter (technical only).
            limit: Max findings (technical only).

        Returns:
            PDF bytes.
        """
        from weasyprint import HTML

        is_compliance = report_type in ("soc2", "iso27001")

        if is_compliance:
            data = self._collect_compliance_data(
                framework="soc2" if report_type == "soc2" else "iso27001",
            )
        else:
            data = self._collect_executive_data()

        charts = self._generate_charts(data, fmt="svg")

        # Load CSS
        css_path = TEMPLATE_DIR / "pdf" / "styles.css"
        css_content = css_path.read_text(encoding="utf-8")

        # Build template context
        tenant_name = self.tenant.name if self.tenant else f"Tenant {self.tenant_id}"
        report_id = uuid.uuid4().hex[:12].upper()
        context: Dict[str, Any] = {
            "css": css_content,
            "tenant_name": tenant_name,
            "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
            "report_id": f"EASM-{report_id}",
            "risk_score": data["risk_score"],
            "risk_grade": data["risk_grade"],
            "total_assets": data["total_assets"],
            "total_findings": data["total_findings"],
            "open_findings": data["open_findings"],
            "severity_counts": data["severity_counts"],
            "top_issues": data["top_issues"],
            "recommendations": data["recommendations"],
            **charts,
        }

        # Choose template
        template_name = f"pdf/{report_type}.html"

        if report_type == "technical":
            findings = self._collect_findings(severity, finding_status, limit)
            # Prepare evidence text for each finding
            for f in findings:
                f["evidence_text"] = _format_evidence(f.get("evidence"))
            context["findings"] = findings
            # Group findings by name+severity → one card per finding type
            # with a compact list of affected assets underneath
            severity_order = ["critical", "high", "medium", "low", "info"]
            grouped: Dict[str, List] = {s: [] for s in severity_order}
            seen_groups: Dict[str, Dict] = {}  # key → group dict
            for f in findings:
                sev = f.get("severity", "info")
                group_key = f"{sev}|{f.get('name', '')}"
                if group_key in seen_groups:
                    seen_groups[group_key]["affected_assets"].append(
                        {
                            "identifier": f["asset_identifier"],
                            "asset_type": f.get("asset_type", "unknown"),
                            "first_seen": f.get("first_seen", "-"),
                            "last_seen": f.get("last_seen", "-"),
                            "status": f.get("status", "open"),
                        }
                    )
                else:
                    group = dict(f)
                    group["affected_assets"] = [
                        {
                            "identifier": f["asset_identifier"],
                            "asset_type": f.get("asset_type", "unknown"),
                            "first_seen": f.get("first_seen", "-"),
                            "last_seen": f.get("last_seen", "-"),
                            "status": f.get("status", "open"),
                        }
                    ]
                    seen_groups[group_key] = group
                    if sev in grouped:
                        grouped[sev].append(group)
            context["severity_order"] = severity_order
            context["grouped_findings"] = grouped

        elif is_compliance:
            # Merge compliance-specific context (categories, mappings, etc.)
            context.update({k: v for k, v in data.items() if k not in context})
            context["asset_counts"] = data["asset_counts"]

        env = Environment(
            loader=FileSystemLoader(str(TEMPLATE_DIR)),
            autoescape=True,
        )
        template = env.get_template(template_name)
        html_string = template.render(**context)

        pdf_bytes = HTML(
            string=html_string,
            base_url=str(TEMPLATE_DIR / "pdf"),
        ).write_pdf(presentational_hints=True)
        return pdf_bytes

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def generate_docx(
        self,
        report_type: str = "executive",
        severity: Optional[str] = None,
        finding_status: Optional[str] = None,
        limit: int = 500,
    ) -> bytes:
        """
        Generate a professional DOCX security report.

        Produces a CrowdStrike/Rapid7-grade document with custom Word styles,
        brand-colored tables, severity-coded finding sections, headers/footers
        with page numbers and confidentiality markings, and full chart
        embedding.

        For compliance report types (``soc2``, ``iso27001``), the DOCX
        output uses the executive template layout. Use PDF export for
        full compliance-specific formatting with TSC/Annex A control mappings.

        Args:
            report_type: ``"executive"``, ``"technical"``, ``"soc2"``,
                         or ``"iso27001"``.
            severity: Optional severity filter (technical only).
            finding_status: Optional status filter (technical only).
            limit: Max findings (technical only).

        Returns:
            DOCX bytes.
        """
        # Compliance types use executive layout for DOCX
        # (full compliance formatting is in PDF templates)
        if report_type in ("soc2", "iso27001"):
            report_type = "executive"
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        # -- Color palette -------------------------------------------------
        BRAND_HEX = "6366F1"
        BRAND = RGBColor(0x63, 0x66, 0xF1)
        BRAND_LIGHT = RGBColor(0x81, 0x8C, 0xF8)
        COVER_BG_HEX = "0F172A"
        COVER_BG = RGBColor(0x0F, 0x17, 0x2A)
        CARD_BG_HEX = "1E293B"
        CARD_BG = RGBColor(0x1E, 0x29, 0x3B)
        WHITE = RGBColor(0xFF, 0xFF, 0xFF)
        NEAR_WHITE = RGBColor(0xF8, 0xFA, 0xFC)
        DARK_TEXT = RGBColor(0x0F, 0x17, 0x2A)
        BODY_TEXT = RGBColor(0x1E, 0x29, 0x3B)
        MUTED_TEXT = RGBColor(0x6B, 0x72, 0x80)
        LIGHT_MUTED = RGBColor(0x94, 0xA3, 0xB8)
        DESC_TEXT = RGBColor(0x47, 0x55, 0x69)
        ROW_BORDER_HEX = "E2E8F0"

        SEV_COLORS = {
            "critical": RGBColor(0xDC, 0x26, 0x26),
            "high": RGBColor(0xEA, 0x58, 0x0C),
            "medium": RGBColor(0xCA, 0x8A, 0x04),
            "low": RGBColor(0x25, 0x63, 0xEB),
            "info": RGBColor(0x6B, 0x72, 0x80),
        }
        SEV_HEX = {
            "critical": "DC2626",
            "high": "EA580C",
            "medium": "CA8A04",
            "low": "2563EB",
            "info": "6B7280",
        }
        SEV_BG_HEX = {
            "critical": "450A0A",
            "high": "431407",
            "medium": "422006",
            "low": "172554",
            "info": "1E293B",
        }
        SEV_TEXT_COLORS = {
            "critical": RGBColor(0xFC, 0xA5, 0xA5),
            "high": RGBColor(0xFD, 0xBA, 0x74),
            "medium": RGBColor(0xFD, 0xE6, 0x8A),
            "low": RGBColor(0x93, 0xC5, 0xFD),
            "info": RGBColor(0x94, 0xA3, 0xB8),
        }
        KPI_COLORS_HEX = {
            "Risk Score": "6366F1",
            "Security Grade": "059669",
            "Total Assets": "2563EB",
            "Open Findings": "DC2626",
        }

        # -- Helper: cell shading ------------------------------------------
        def _shade_cell(cell, hex_color: str):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # -- Helper: set cell width ----------------------------------------
        def _set_cell_width(cell, width_inches: float):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_inches * 1440)}" w:type="dxa"/>')
            tc_pr.append(tc_w)

        # -- Helper: set cell vertical alignment ---------------------------
        def _set_cell_valign(cell, val: str = "center"):
            tc_pr = cell._tc.get_or_add_tcPr()
            va = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
            tc_pr.append(va)

        # -- Helper: set cell margins (tight padding) ----------------------
        def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = parse_xml(
                f"<w:tcMar {nsdecls('w')}>"
                f'  <w:top w:w="{top}" w:type="dxa"/>'
                f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
                f'  <w:left w:w="{left}" w:type="dxa"/>'
                f'  <w:right w:w="{right}" w:type="dxa"/>'
                f"</w:tcMar>"
            )
            tc_pr.append(margins)

        # -- Helper: remove all table borders ------------------------------
        def _remove_table_borders(table):
            tbl = table._tbl
            tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            borders = parse_xml(
                f"<w:tblBorders {nsdecls('w')}>"
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f"</w:tblBorders>"
            )
            # Remove existing borders element if present
            for existing in tbl_pr.findall(qn("w:tblBorders")):
                tbl_pr.remove(existing)
            tbl_pr.append(borders)

        # -- Helper: set clean table borders (light grid) ------------------
        def _set_table_light_borders(table, color_hex: str = ROW_BORDER_HEX):
            tbl = table._tbl
            tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f"<w:tblPr {nsdecls('w')}/>")
            borders = parse_xml(
                f"<w:tblBorders {nsdecls('w')}>"
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>'
                f"</w:tblBorders>"
            )
            for existing in tbl_pr.findall(qn("w:tblBorders")):
                tbl_pr.remove(existing)
            tbl_pr.append(borders)

        # -- Helper: style a dark header row (dark bg, light text) -----------
        def _style_header_row(table, headers, col_widths=None):
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                cell = hdr_cells[i]
                cell.text = ""
                _shade_cell(cell, CARD_BG_HEX)
                _set_cell_valign(cell)
                _set_cell_margins(cell, top=50, bottom=50, left=80, right=80)
                if col_widths and i < len(col_widths):
                    _set_cell_width(cell, col_widths[i])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(h.upper())
                run.bold = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                run.font.name = "Inter"
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

        # -- Helper: add a clean data row (no alternating shading) ----------
        def _add_data_row(table, values, row_idx, font_size=Pt(8), bold_cols=None, color_map=None):
            row_cells = table.add_row().cells
            for i, val in enumerate(values):
                cell = row_cells[i]
                cell.text = ""
                _set_cell_valign(cell)
                _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(val))
                run.font.size = font_size
                run.font.name = "Inter"
                run.font.color.rgb = BODY_TEXT
                if bold_cols and i in bold_cols:
                    run.bold = True
                if color_map and i in color_map:
                    run.font.color.rgb = color_map[i]
            return row_cells

        # -- Helper: add paragraph with controlled spacing -----------------
        def _add_spaced_para(doc, text="", before=0, after=6, size=Pt(10), color=BODY_TEXT, bold=False, alignment=None):
            p = doc.add_paragraph()
            if alignment is not None:
                p.alignment = alignment
            p.paragraph_format.space_before = Pt(before)
            p.paragraph_format.space_after = Pt(after)
            if text:
                run = p.add_run(text)
                run.font.size = size
                run.font.color.rgb = color
                run.font.name = "Inter"
                run.bold = bold
            return p

        # -- Helper: add a section divider line ----------------------------
        def _add_divider(doc, color_hex=BRAND_HEX, width_pct=15):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f"<w:pBdr {nsdecls('w')}>"
                f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{color_hex}"/>'
                f"</w:pBdr>"
            )
            pPr.append(pBdr)

        # ==================================================================
        # Collect data
        # ==================================================================
        data = self._collect_executive_data()
        tenant_name = self.tenant.name if self.tenant else f"Tenant {self.tenant_id}"
        now_display = datetime.now(timezone.utc).strftime("%B %d, %Y  %H:%M UTC")
        now_short = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        report_id = f"EASM-{uuid.uuid4().hex[:12].upper()}"
        subtitle_text = "Executive Summary" if report_type == "executive" else "Technical Assessment Report"

        doc = Document()

        # ==================================================================
        # GLOBAL STYLE CONFIGURATION
        # ==================================================================
        # -- Normal body style
        style_normal = doc.styles["Normal"]
        style_normal.font.name = "Inter"
        style_normal.font.size = Pt(10)
        style_normal.font.color.rgb = BODY_TEXT
        style_normal.paragraph_format.space_after = Pt(6)
        style_normal.paragraph_format.line_spacing = 1.15

        # -- Heading 1: section headers with brand color, no border
        h1_style = doc.styles["Heading 1"]
        h1_style.font.name = "Inter"
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        h1_style.paragraph_format.space_before = Pt(28)
        h1_style.paragraph_format.space_after = Pt(10)
        h1_style.paragraph_format.keep_with_next = True

        # -- Heading 2: subsection headers
        h2_style = doc.styles["Heading 2"]
        h2_style.font.name = "Inter"
        h2_style.font.size = Pt(13)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        h2_style.paragraph_format.space_before = Pt(18)
        h2_style.paragraph_format.space_after = Pt(6)
        h2_style.paragraph_format.keep_with_next = True

        # -- Heading 3: chart/sub-subsection labels
        h3_style = doc.styles["Heading 3"]
        h3_style.font.name = "Inter"
        h3_style.font.size = Pt(11)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
        h3_style.paragraph_format.space_before = Pt(12)
        h3_style.paragraph_format.space_after = Pt(4)
        h3_style.paragraph_format.keep_with_next = True

        # ==================================================================
        # PAGE LAYOUT: narrow margins for more content width
        # ==================================================================
        for section in doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

        # ==================================================================
        # HEADERS AND FOOTERS
        # ==================================================================
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # -- Default header (pages 2+): tenant name on left, report type right
        header = section.header
        header.is_linked_to_previous = False
        h_para = header.paragraphs[0]
        h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_para.paragraph_format.space_after = Pt(2)
        run_left = h_para.add_run(f"EASM {subtitle_text}")
        run_left.font.size = Pt(7.5)
        run_left.font.color.rgb = LIGHT_MUTED
        run_left.font.name = "Inter"
        h_para.add_run("  |  ").font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        run_tenant = h_para.add_run(tenant_name)
        run_tenant.font.size = Pt(7.5)
        run_tenant.font.color.rgb = LIGHT_MUTED
        run_tenant.font.name = "Inter"
        # Header bottom border
        h_pPr = h_para._p.get_or_add_pPr()
        h_bdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="E5E7EB"/></w:pBdr>'
        )
        h_pPr.append(h_bdr)

        # -- First page header: empty (cover page)
        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        first_header.paragraphs[0].text = ""

        # -- Default footer: page numbers center, confidential right
        footer = section.footer
        footer.is_linked_to_previous = False
        f_para = footer.paragraphs[0]
        f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_para.paragraph_format.space_before = Pt(4)
        # Top border on footer
        f_pPr = f_para._p.get_or_add_pPr()
        f_bdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>  <w:top w:val="single" w:sz="4" w:space="6" w:color="E5E7EB"/></w:pBdr>'
        )
        f_pPr.append(f_bdr)

        run_conf = f_para.add_run("CONFIDENTIAL")
        run_conf.font.size = Pt(6.5)
        run_conf.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        run_conf.font.name = "Inter"
        run_conf.font.all_caps = True
        f_para.add_run("    ").font.size = Pt(6.5)

        run_pre = f_para.add_run("Page ")
        run_pre.font.size = Pt(7.5)
        run_pre.font.color.rgb = LIGHT_MUTED
        run_pre.font.name = "Inter"
        # Insert PAGE field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fld_instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_page = f_para.add_run()
        run_page.font.size = Pt(7.5)
        run_page.font.color.rgb = LIGHT_MUTED
        run_page._r.append(fld_char_begin)
        run_page._r.append(fld_instr)
        run_page._r.append(fld_char_end)

        run_of = f_para.add_run(" of ")
        run_of.font.size = Pt(7.5)
        run_of.font.color.rgb = LIGHT_MUTED
        run_of.font.name = "Inter"
        # Insert NUMPAGES field
        fld_char_begin2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fld_instr2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        fld_char_end2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_numpages = f_para.add_run()
        run_numpages.font.size = Pt(7.5)
        run_numpages.font.color.rgb = LIGHT_MUTED
        run_numpages._r.append(fld_char_begin2)
        run_numpages._r.append(fld_instr2)
        run_numpages._r.append(fld_char_end2)

        f_para.add_run("    ").font.size = Pt(6.5)
        run_date = f_para.add_run(now_short)
        run_date.font.size = Pt(6.5)
        run_date.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        run_date.font.name = "Inter"

        # -- First page footer: empty (cover page)
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        first_footer.paragraphs[0].text = ""

        # ==================================================================
        # COVER PAGE (dark theme, PushingTorPod identity)
        # ==================================================================
        # Full-page dark cover via a single-cell table
        cover_tbl = doc.add_table(rows=1, cols=1)
        cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(cover_tbl)
        cover_cell = cover_tbl.rows[0].cells[0]
        _shade_cell(cover_cell, COVER_BG_HEX)
        _set_cell_margins(cover_cell, top=600, bottom=400, left=400, right=400)

        # Brand label
        brand_p = cover_cell.paragraphs[0]
        brand_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        brand_p.paragraph_format.space_before = Pt(80)
        brand_p.paragraph_format.space_after = Pt(4)
        brand_run = brand_p.add_run("PUSHINGTORPOD")
        brand_run.font.size = Pt(7.5)
        brand_run.font.color.rgb = BRAND_LIGHT
        brand_run.font.name = "Inter"
        brand_run.bold = True

        # Brand sub-label
        brand_sub_p = cover_cell.add_paragraph()
        brand_sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        brand_sub_p.paragraph_format.space_before = Pt(0)
        brand_sub_p.paragraph_format.space_after = Pt(12)
        brand_sub_run = brand_sub_p.add_run("External Attack Surface Management")
        brand_sub_run.font.size = Pt(6.5)
        brand_sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        brand_sub_run.font.name = "Inter"

        # Title
        title_p = cover_cell.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(4)
        run_t = title_p.add_run("Security Assessment")
        run_t.font.size = Pt(36)
        run_t.font.color.rgb = NEAR_WHITE
        run_t.bold = True
        run_t.font.name = "Inter"

        # Subtitle (report type)
        sub_p = cover_cell.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_p.paragraph_format.space_before = Pt(4)
        sub_p.paragraph_format.space_after = Pt(16)
        run_sub = sub_p.add_run(subtitle_text)
        run_sub.font.size = Pt(14)
        run_sub.font.color.rgb = BRAND_LIGHT
        run_sub.font.name = "Inter"

        # Thin separator line
        div_p = cover_cell.add_paragraph()
        div_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        div_p.paragraph_format.space_before = Pt(4)
        div_p.paragraph_format.space_after = Pt(16)
        pPr = div_p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="{BRAND_HEX}"/></w:pBdr>'
        )
        pPr.append(pBdr)

        # Risk score prominent display
        score_p = cover_cell.add_paragraph()
        score_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        score_p.paragraph_format.space_before = Pt(8)
        score_p.paragraph_format.space_after = Pt(2)
        score_run = score_p.add_run(f"{data['risk_score']:.0f}")
        score_run.font.size = Pt(48)
        score_run.font.color.rgb = NEAR_WHITE
        score_run.bold = True
        score_run.font.name = "Inter"
        score_label_p = cover_cell.add_paragraph()
        score_label_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        score_label_p.paragraph_format.space_before = Pt(0)
        score_label_p.paragraph_format.space_after = Pt(20)
        score_label_run = score_label_p.add_run("RISK SCORE / 100")
        score_label_run.font.size = Pt(8)
        score_label_run.font.color.rgb = LIGHT_MUTED
        score_label_run.font.name = "Inter"
        score_label_run.bold = True

        # Metadata lines on dark background
        meta_items = [
            ("Organization", tenant_name),
            ("Date", now_display),
            ("Report ID", report_id),
            ("Classification", "CONFIDENTIAL"),
        ]
        for label, value in meta_items:
            mp = cover_cell.add_paragraph()
            mp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            mp.paragraph_format.space_before = Pt(2)
            mp.paragraph_format.space_after = Pt(2)
            lbl_run = mp.add_run(f"{label}:  ")
            lbl_run.font.size = Pt(9)
            lbl_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            lbl_run.font.name = "Inter"
            val_run = mp.add_run(value)
            val_run.font.size = Pt(9)
            val_run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
            val_run.font.name = "Inter"
            val_run.bold = True

        # Spacer
        sp = cover_cell.add_paragraph()
        sp.paragraph_format.space_before = Pt(20)
        sp.paragraph_format.space_after = Pt(0)

        # Confidentiality notice
        conf_p = cover_cell.add_paragraph()
        conf_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        conf_p.paragraph_format.space_before = Pt(0)
        conf_p.paragraph_format.space_after = Pt(0)
        run_c = conf_p.add_run(
            "This document contains confidential security assessment data. "
            "Distribution is restricted to authorized recipients only."
        )
        run_c.font.size = Pt(7.5)
        run_c.font.color.rgb = MUTED_TEXT
        run_c.font.name = "Inter"
        run_c.italic = True

        doc.add_page_break()

        # ==================================================================
        # TABLE OF CONTENTS (placeholder)
        # ==================================================================
        doc.add_heading("Table of Contents", level=1)
        toc_items = [
            ("1.", "Executive Summary"),
            ("2.", "Risk Overview"),
            ("   2.1", "Risk Score and Severity Breakdown"),
            ("   2.2", "Asset Distribution and Trend"),
            ("3.", "Top 10 Issues"),
            ("4.", "Recommendations"),
        ]
        if report_type == "technical":
            toc_items.append(("5.", "Detailed Findings"))

        for num, label in toc_items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            is_indent = num.startswith("   ")
            if is_indent:
                p.paragraph_format.left_indent = Cm(1.0)
            run_num = p.add_run(f"{num.strip()}  ")
            run_num.font.size = Pt(10) if not is_indent else Pt(9.5)
            run_num.font.color.rgb = BRAND
            run_num.bold = True
            run_num.font.name = "Inter"
            run_lbl = p.add_run(label)
            run_lbl.font.size = Pt(10) if not is_indent else Pt(9.5)
            run_lbl.font.color.rgb = BODY_TEXT if not is_indent else DESC_TEXT
            run_lbl.font.name = "Inter"

        doc.add_page_break()

        # ==================================================================
        # 1. EXECUTIVE SUMMARY
        # ==================================================================
        doc.add_heading("1. Executive Summary", level=1)

        # Narrative paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(12)
        r1 = p.add_run("This report provides an assessment of the external attack surface for ")
        r1.font.size = Pt(10)
        r1.font.name = "Inter"
        r2 = p.add_run(tenant_name)
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.name = "Inter"
        r3 = p.add_run(
            f". The organization manages {data['total_assets']:,} monitored assets "
            f"with {data['open_findings']:,} open "
            f"finding{'s' if data['open_findings'] != 1 else ''}. "
            f"The overall risk score is "
        )
        r3.font.size = Pt(10)
        r3.font.name = "Inter"
        r4 = p.add_run(f"{data['risk_score']:.0f}/100")
        r4.bold = True
        r4.font.size = Pt(10)
        r4.font.name = "Inter"
        r5 = p.add_run(f" (Grade ")
        r5.font.size = Pt(10)
        r5.font.name = "Inter"
        r6 = p.add_run(f"{data['risk_grade']}")
        r6.bold = True
        r6.font.size = Pt(10)
        r6.font.name = "Inter"
        r7 = p.add_run(").")
        r7.font.size = Pt(10)
        r7.font.name = "Inter"

        # -- KPI Summary Cards (4-column table, 2 rows) --------------------
        doc.add_heading("Key Performance Indicators", level=3)

        kpi_data = [
            ("Risk Score", f"{data['risk_score']:.0f}", "/100"),
            ("Security Grade", data["risk_grade"], ""),
            ("Total Assets", f"{data['total_assets']:,}", "monitored"),
            ("Open Findings", f"{data['open_findings']:,}", "active"),
        ]
        kpi_tbl = doc.add_table(rows=2, cols=4)
        kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(kpi_tbl)

        for i, (kpi_label, kpi_value, kpi_suffix) in enumerate(kpi_data):
            kpi_color = list(KPI_COLORS_HEX.values())[i]
            kpi_rgb = RGBColor(
                int(kpi_color[0:2], 16),
                int(kpi_color[2:4], 16),
                int(kpi_color[4:6], 16),
            )

            # Label cell (dark background, muted text)
            hdr_cell = kpi_tbl.rows[0].cells[i]
            _shade_cell(hdr_cell, CARD_BG_HEX)
            _set_cell_valign(hdr_cell)
            _set_cell_margins(hdr_cell, top=50, bottom=30, left=60, right=60)
            hdr_cell.text = ""
            hp = hdr_cell.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hp.paragraph_format.space_before = Pt(0)
            hp.paragraph_format.space_after = Pt(0)
            hr = hp.add_run(kpi_label.upper())
            hr.bold = True
            hr.font.size = Pt(7)
            hr.font.color.rgb = LIGHT_MUTED
            hr.font.name = "Inter"

            # Value cell (dark background, near-white value)
            val_cell = kpi_tbl.rows[1].cells[i]
            _shade_cell(val_cell, CARD_BG_HEX)
            _set_cell_valign(val_cell)
            _set_cell_margins(val_cell, top=30, bottom=70, left=60, right=60)
            val_cell.text = ""
            vp = val_cell.paragraphs[0]
            vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vp.paragraph_format.space_before = Pt(0)
            vp.paragraph_format.space_after = Pt(0)
            vr = vp.add_run(kpi_value)
            vr.bold = True
            vr.font.size = Pt(22)
            vr.font.color.rgb = NEAR_WHITE
            vr.font.name = "Inter"
            if kpi_suffix:
                sr = vp.add_run(f" {kpi_suffix}")
                sr.font.size = Pt(8)
                sr.font.color.rgb = LIGHT_MUTED
                sr.font.name = "Inter"

        _add_spaced_para(doc, before=10, after=6)

        # -- Severity Breakdown Table --------------------------------------
        doc.add_heading("Severity Breakdown", level=3)

        total_findings = data["total_findings"] or 1
        sev_tbl = doc.add_table(rows=1, cols=4)
        sev_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_light_borders(sev_tbl)
        _style_header_row(sev_tbl, ["Severity", "Count", "% of Total", "Risk Level"], col_widths=[2.0, 1.2, 1.5, 1.8])

        risk_labels = {
            "critical": "Immediate Action",
            "high": "High Priority",
            "medium": "Planned Fix",
            "low": "Acceptable Risk",
            "info": "Informational",
        }
        for row_idx, sev in enumerate(["critical", "high", "medium", "low", "info"]):
            count = data["severity_counts"].get(sev, 0)
            pct = f"{count / total_findings * 100:.1f}%" if total_findings > 0 else "0.0%"
            row_cells = sev_tbl.add_row().cells

            for ci, cell in enumerate(row_cells):
                _set_cell_valign(cell)
                _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)

            # Severity name with left color indicator
            row_cells[0].text = ""
            sp = row_cells[0].paragraphs[0]
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(0)
            # Color indicator block character
            indicator_run = sp.add_run("\u2588 ")
            indicator_run.font.size = Pt(9)
            indicator_run.font.color.rgb = SEV_COLORS[sev]
            sev_run = sp.add_run(sev.upper())
            sev_run.bold = True
            sev_run.font.size = Pt(8.5)
            sev_run.font.color.rgb = SEV_COLORS[sev]
            sev_run.font.name = "Inter"

            # Count
            row_cells[1].text = ""
            cp = row_cells[1].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            cr = cp.add_run(str(count))
            cr.bold = True
            cr.font.size = Pt(9)
            cr.font.name = "Inter"
            cr.font.color.rgb = BODY_TEXT

            # Percentage
            row_cells[2].text = ""
            pp = row_cells[2].paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(0)
            pp.paragraph_format.space_after = Pt(0)
            pr = pp.add_run(pct)
            pr.font.size = Pt(9)
            pr.font.name = "Inter"
            pr.font.color.rgb = BODY_TEXT

            # Risk level label
            row_cells[3].text = ""
            rp = row_cells[3].paragraphs[0]
            rp.paragraph_format.space_before = Pt(0)
            rp.paragraph_format.space_after = Pt(0)
            rr = rp.add_run(risk_labels[sev])
            rr.font.size = Pt(8)
            rr.font.name = "Inter"
            rr.italic = True
            rr.font.color.rgb = DESC_TEXT

        # Totals row (light top border, no background)
        total_cells = sev_tbl.add_row().cells
        for cell in total_cells:
            _set_cell_valign(cell)
            _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
            # Top border for totals separation
            tc_pr = cell._tc.get_or_add_tcPr()
            top_bdr = parse_xml(
                f"<w:tcBorders {nsdecls('w')}>"
                f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/>'
                f"</w:tcBorders>"
            )
            tc_pr.append(top_bdr)

        total_cells[0].text = ""
        tp = total_cells[0].paragraphs[0]
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(0)
        tr = tp.add_run("TOTAL")
        tr.bold = True
        tr.font.size = Pt(8.5)
        tr.font.name = "Inter"
        tr.font.color.rgb = DARK_TEXT

        total_cells[1].text = ""
        tc_p = total_cells[1].paragraphs[0]
        tc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_p.paragraph_format.space_before = Pt(0)
        tc_p.paragraph_format.space_after = Pt(0)
        tc_r = tc_p.add_run(str(data["total_findings"]))
        tc_r.bold = True
        tc_r.font.size = Pt(9)
        tc_r.font.name = "Inter"
        tc_r.font.color.rgb = DARK_TEXT

        total_cells[2].text = ""
        tp2 = total_cells[2].paragraphs[0]
        tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp2.paragraph_format.space_before = Pt(0)
        tp2.paragraph_format.space_after = Pt(0)
        tr2 = tp2.add_run("100.0%")
        tr2.bold = True
        tr2.font.size = Pt(9)
        tr2.font.name = "Inter"
        tr2.font.color.rgb = DARK_TEXT

        total_cells[3].text = ""
        tp3 = total_cells[3].paragraphs[0]
        tp3.paragraph_format.space_before = Pt(0)
        tp3.paragraph_format.space_after = Pt(0)

        _add_spaced_para(doc, before=6, after=2)

        # ==================================================================
        # 2. RISK OVERVIEW (Charts)
        # ==================================================================
        doc.add_page_break()
        doc.add_heading("2. Risk Overview", level=1)

        _add_spaced_para(
            doc,
            text="The following charts illustrate the current risk posture, "
            "severity distribution across findings, asset composition, "
            "and 30-day risk score trend.",
            before=2,
            after=10,
            size=Pt(9.5),
            color=DESC_TEXT,
        )

        chart_images = self._generate_chart_images_png(data)
        chart_section_labels = {
            "Risk Score Gauge": "2.1 Risk Score Gauge",
            "Findings by Severity": "2.1 Findings by Severity",
            "Asset Distribution": "2.2 Asset Distribution by Type",
            "Risk Score Trend (30 days)": "2.2 Risk Score Trend (30 Days)",
        }

        chart_count = 0
        for label, img_bytes in chart_images.items():
            if not img_bytes:
                continue
            chart_count += 1
            section_label = chart_section_labels.get(label, label)
            doc.add_heading(section_label, level=3)
            stream = io.BytesIO(img_bytes)
            doc.add_picture(stream, width=Inches(5.0))
            # Center the image
            last_para = doc.paragraphs[-1]
            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_para.paragraph_format.space_after = Pt(12)

        if chart_count == 0:
            _add_spaced_para(
                doc,
                text="No chart data available for the current period.",
                before=4,
                after=8,
                size=Pt(9.5),
                color=MUTED_TEXT,
                bold=False,
            )

        # ==================================================================
        # 3. TOP 10 ISSUES
        # ==================================================================
        doc.add_page_break()
        doc.add_heading("3. Top 10 Issues", level=1)

        _add_spaced_para(
            doc,
            text="The following table lists the highest-priority open findings "
            "ranked by a composite score of severity weight and CVSS rating.",
            before=2,
            after=10,
            size=Pt(9.5),
            color=DESC_TEXT,
        )

        if data["top_issues"]:
            col_widths = [0.35, 2.4, 0.9, 0.6, 1.7, 0.85]
            issues_tbl = doc.add_table(rows=1, cols=6)
            issues_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_light_borders(issues_tbl)
            _style_header_row(
                issues_tbl,
                ["#", "Finding", "Severity", "CVSS", "Affected Asset", "First Seen"],
                col_widths=col_widths,
            )

            for idx, issue in enumerate(data["top_issues"], 1):
                sev = issue.get("severity", "info")
                cvss_text = f"{issue['cvss_score']:.1f}" if issue.get("cvss_score") else "--"
                sev_color = SEV_COLORS.get(sev, BODY_TEXT)
                row_cells = _add_data_row(
                    issues_tbl,
                    [
                        str(idx),
                        issue["name"],
                        sev.upper(),
                        cvss_text,
                        issue["asset_identifier"],
                        issue.get("first_seen", "--"),
                    ],
                    row_idx=idx,
                    font_size=Pt(8),
                    bold_cols={2},
                    color_map={2: sev_color},
                )
                # Severity cell dark background + light text
                _shade_cell(row_cells[2], SEV_BG_HEX.get(sev, "1E293B"))
                for run in row_cells[2].paragraphs[0].runs:
                    run.font.color.rgb = SEV_TEXT_COLORS.get(sev, LIGHT_MUTED)
        else:
            _add_spaced_para(
                doc,
                text="No open issues found. The external attack surface appears clean.",
                before=4,
                after=8,
                size=Pt(10),
                color=MUTED_TEXT,
            )

        _add_spaced_para(doc, before=6, after=2)

        # ==================================================================
        # 4. RECOMMENDATIONS
        # ==================================================================
        doc.add_heading("4. Recommendations", level=1)

        _add_spaced_para(
            doc,
            text="Based on the current assessment, the following actions are recommended in order of priority.",
            before=2,
            after=10,
            size=Pt(9.5),
            color=DESC_TEXT,
        )

        for rec in data["recommendations"]:
            # Recommendation as a single-row, 2-column table (priority badge | text)
            rec_tbl = doc.add_table(rows=1, cols=2)
            rec_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            _remove_table_borders(rec_tbl)

            # Priority badge cell (dark, brand indigo text)
            badge_cell = rec_tbl.rows[0].cells[0]
            _set_cell_width(badge_cell, 0.45)
            _shade_cell(badge_cell, "0F172A")
            _set_cell_valign(badge_cell, "center")
            _set_cell_margins(badge_cell, top=50, bottom=50, left=50, right=50)
            badge_cell.text = ""
            bp = badge_cell.paragraphs[0]
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(0)
            br_run = bp.add_run(str(rec["priority"]))
            br_run.bold = True
            br_run.font.size = Pt(11)
            br_run.font.color.rgb = BRAND_LIGHT
            br_run.font.name = "Inter"

            # Content cell
            content_cell = rec_tbl.rows[0].cells[1]
            _set_cell_width(content_cell, 6.1)
            _shade_cell(content_cell, "F9FAFB")
            _set_cell_valign(content_cell, "center")
            _set_cell_margins(content_cell, top=50, bottom=50, left=100, right=80)
            content_cell.text = ""

            # Title
            title_p = content_cell.paragraphs[0]
            title_p.paragraph_format.space_before = Pt(0)
            title_p.paragraph_format.space_after = Pt(2)
            title_run = title_p.add_run(rec["title"])
            title_run.bold = True
            title_run.font.size = Pt(10)
            title_run.font.color.rgb = DARK_TEXT
            title_run.font.name = "Inter"

            # Description
            desc_p = content_cell.add_paragraph()
            desc_p.paragraph_format.space_before = Pt(0)
            desc_p.paragraph_format.space_after = Pt(0)
            desc_run = desc_p.add_run(rec["description"])
            desc_run.font.size = Pt(9)
            desc_run.font.color.rgb = DESC_TEXT
            desc_run.font.name = "Inter"

            # Affected count tag (if nonzero)
            if rec.get("affected_count", 0) > 0:
                tag_p = content_cell.add_paragraph()
                tag_p.paragraph_format.space_before = Pt(3)
                tag_p.paragraph_format.space_after = Pt(0)
                tag_run = tag_p.add_run(
                    f"{rec['affected_count']} finding{'s' if rec['affected_count'] != 1 else ''} affected"
                )
                tag_run.font.size = Pt(7.5)
                tag_run.font.color.rgb = BRAND
                tag_run.bold = True
                tag_run.font.name = "Inter"

            # Spacer after each recommendation
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after = Pt(2)

        # ==================================================================
        # 5. TECHNICAL FINDINGS (technical report only)
        # ==================================================================
        if report_type == "technical":
            doc.add_page_break()
            findings = self._collect_findings(severity, finding_status, limit)

            doc.add_heading(f"5. Detailed Findings ({len(findings)})", level=1)

            _add_spaced_para(
                doc,
                text=(
                    "This section provides a detailed breakdown of all findings, "
                    "grouped by severity level. Each finding includes the affected "
                    "asset, CVSS score, occurrence count, and recommended remediation."
                ),
                before=2,
                after=10,
                size=Pt(9.5),
                color=DESC_TEXT,
            )

            if findings:
                # Group by severity
                severity_order = ["critical", "high", "medium", "low", "info"]
                grouped: Dict[str, List] = {s: [] for s in severity_order}
                for f in findings:
                    sev = f.get("severity", "info")
                    if sev in grouped:
                        grouped[sev].append(f)

                for sev in severity_order:
                    sev_findings = grouped.get(sev, [])
                    if not sev_findings:
                        continue

                    # -- Severity group header (dark banner) ----------------
                    SEV_BANNER_BG = {
                        "critical": "7F1D1D",
                        "high": "7C2D12",
                        "medium": "713F12",
                        "low": "1E3A5F",
                        "info": "334155",
                    }
                    banner_tbl = doc.add_table(rows=1, cols=1)
                    banner_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                    _remove_table_borders(banner_tbl)
                    banner_cell = banner_tbl.rows[0].cells[0]
                    _shade_cell(banner_cell, SEV_BANNER_BG.get(sev, "334155"))
                    _set_cell_margins(banner_cell, top=50, bottom=50, left=100, right=80)

                    banner_cell.text = ""
                    bp = banner_cell.paragraphs[0]
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after = Pt(0)
                    sev_icon = {
                        "critical": "\u26a0",
                        "high": "\u26a0",
                        "medium": "\u25b2",
                        "low": "\u25cf",
                        "info": "\u2139",
                    }
                    icon_run = bp.add_run(f"{sev_icon.get(sev, '')} ")
                    icon_run.font.size = Pt(11)
                    icon_run.font.color.rgb = WHITE
                    sev_title_run = bp.add_run(f"{sev.upper()} SEVERITY")
                    sev_title_run.bold = True
                    sev_title_run.font.size = Pt(12)
                    sev_title_run.font.color.rgb = WHITE
                    sev_title_run.font.name = "Inter"
                    count_run = bp.add_run(f"  --  {len(sev_findings)} finding{'s' if len(sev_findings) != 1 else ''}")
                    count_run.font.size = Pt(10)
                    count_run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
                    count_run.font.name = "Inter"

                    _add_spaced_para(doc, before=4, after=2)

                    # -- Findings table for this severity group ------------
                    col_widths_tech = [0.4, 1.9, 0.55, 1.5, 0.6, 0.45, 1.5]
                    t = doc.add_table(rows=1, cols=7)
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    _set_table_light_borders(t)
                    _style_header_row(
                        t,
                        ["ID", "Finding", "CVSS", "Asset", "Status", "Seen", "Remediation"],
                        col_widths=col_widths_tech,
                    )

                    for fidx, f in enumerate(sev_findings):
                        cvss_text = f"{f['cvss_score']:.1f}" if f.get("cvss_score") else "--"
                        occ = f.get("occurrence_count", 1) or 1
                        occ_text = f"{occ}x" if occ > 1 else "1x"
                        remediation_text = f.get("remediation") or "--"
                        # Truncate long remediation for table readability
                        if len(remediation_text) > 80:
                            remediation_text = remediation_text[:77] + "..."

                        row_cells = t.add_row().cells

                        values = [
                            str(f["id"]),
                            f["name"],
                            cvss_text,
                            f["asset_identifier"],
                            f["status"].upper(),
                            occ_text,
                            remediation_text,
                        ]

                        for ci, cell in enumerate(row_cells):
                            _set_cell_valign(cell)
                            _set_cell_margins(cell, top=35, bottom=35, left=60, right=60)
                            cell.text = ""
                            cp = cell.paragraphs[0]
                            cp.paragraph_format.space_before = Pt(0)
                            cp.paragraph_format.space_after = Pt(0)
                            run = cp.add_run(values[ci])
                            run.font.size = Pt(7.5)
                            run.font.name = "Inter"
                            run.font.color.rgb = BODY_TEXT

                        # Bold the finding name (col 1)
                        for run in row_cells[1].paragraphs[0].runs:
                            run.bold = True

                        # Color the status cell
                        status_val = f["status"].lower()
                        status_colors = {
                            "open": RGBColor(0xDC, 0x26, 0x26),
                            "fixed": RGBColor(0x16, 0x65, 0x34),
                            "suppressed": RGBColor(0x6B, 0x72, 0x80),
                        }
                        for run in row_cells[4].paragraphs[0].runs:
                            run.bold = True
                            run.font.color.rgb = status_colors.get(status_val, BODY_TEXT)

                        # Highlight occurrence count > 1
                        if occ > 1:
                            _shade_cell(row_cells[5], SEV_BG_HEX.get(sev, "1E293B"))
                            for run in row_cells[5].paragraphs[0].runs:
                                run.bold = True
                                run.font.color.rgb = SEV_TEXT_COLORS.get(sev, LIGHT_MUTED)
                            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    _add_spaced_para(doc, before=8, after=4)

            else:
                _add_spaced_para(
                    doc,
                    text="No findings match the current filters.",
                    before=4,
                    after=8,
                    size=Pt(10),
                    color=MUTED_TEXT,
                )

        # ==================================================================
        # FINAL PAGE: Disclaimer / End of Report
        # ==================================================================
        doc.add_page_break()

        _add_spaced_para(doc, before=40, after=0)

        end_p = doc.add_paragraph()
        end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_p.paragraph_format.space_before = Pt(20)
        end_p.paragraph_format.space_after = Pt(8)
        end_run = end_p.add_run("End of Report")
        end_run.font.size = Pt(14)
        end_run.font.color.rgb = BRAND_LIGHT
        end_run.bold = True
        end_run.font.name = "Inter"

        _add_divider(doc)

        disclaimer_lines = [
            "This report was generated automatically by the EASM platform. "
            "Findings are based on automated scanning and may require manual "
            "verification before remediation actions are taken.",
            "",
            f"Report ID: {report_id}",
            f"Generated: {now_display}",
            f"Classification: CONFIDENTIAL",
        ]
        for line in disclaimer_lines:
            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dp.paragraph_format.space_before = Pt(1)
            dp.paragraph_format.space_after = Pt(1)
            if line:
                dr = dp.add_run(line)
                dr.font.size = Pt(8)
                dr.font.color.rgb = LIGHT_MUTED
                dr.font.name = "Inter"
                dr.italic = True

        # ==================================================================
        # Serialize and return
        # ==================================================================
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def _generate_chart_images_png(self, data: Dict[str, Any]) -> Dict[str, Optional[bytes]]:
        """Generate PNG chart images for DOCX embedding."""
        result: Dict[str, Optional[bytes]] = {}

        result["Risk Score Gauge"] = generate_risk_gauge(data["risk_score"], data["risk_grade"], fmt="png")
        result["Findings by Severity"] = generate_severity_chart(data["severity_counts"], fmt="png")
        result["Asset Distribution"] = generate_asset_chart(data["asset_counts"], fmt="png")
        result["Risk Score Trend (30 days)"] = generate_trend_chart(data["score_trend"], fmt="png")

        return result
